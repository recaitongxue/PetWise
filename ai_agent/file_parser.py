"""
File Parser Module
Supports parsing of MD, DOCX, DOC, and PDF files for knowledge base import
"""
import os
import re
import json
from typing import Dict, List, Any, Optional

from config import Config
from utils import setup_logging

logger = setup_logging(Config.LOG_LEVEL)

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    logger.warning("markdown library not installed, MD files will be read as plain text")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx library not installed, DOCX files will not be supported")

try:
    import fitz
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyMuPDF library not installed, PDF files will not be supported")

class FileParser:
    """File parser for various document formats"""
    
    SUPPORTED_EXTENSIONS = ['.md', '.docx', '.doc', '.pdf', '.txt']
    
    @classmethod
    def get_supported_formats(cls) -> List[str]:
        """Get list of supported file formats"""
        formats = []
        if MARKDOWN_AVAILABLE:
            formats.append('.md')
        if DOCX_AVAILABLE:
            formats.append('.docx')
        if PDF_AVAILABLE:
            formats.append('.pdf')
        formats.append('.txt')
        return formats
    
    @classmethod
    def parse_file(cls, file_path: str) -> Dict[str, Any]:
        """
        Parse a file and extract knowledge entries
        
        Args:
            file_path: Path to the file to parse
            
        Returns:
            Dictionary with title, content, and extracted entries
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {ext}")
        
        logger.info(f"Parsing file: {file_path}")
        
        try:
            if ext == '.md':
                return cls._parse_md(file_path)
            elif ext == '.docx':
                return cls._parse_docx(file_path)
            elif ext == '.pdf':
                return cls._parse_pdf(file_path)
            elif ext == '.txt':
                return cls._parse_txt(file_path)
            else:
                raise ValueError(f"Unsupported format: {ext}")
        except Exception as e:
            logger.error(f"Failed to parse file {file_path}: {e}")
            raise
    
    @classmethod
    def _parse_md(cls, file_path: str) -> Dict[str, Any]:
        """Parse Markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            title = cls._extract_title_from_md(content)
            entries = cls._extract_entries_from_md(content)
            
            result = {
                "title": title,
                "content": content,
                "entries": entries,
                "format": "markdown",
                "file_name": os.path.basename(file_path)
            }
            
            logger.info(f"Parsed MD file: {title}, {len(entries)} entries")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing MD file: {e}")
            raise
    
    @classmethod
    def _extract_title_from_md(cls, content: str) -> str:
        """Extract title from Markdown content"""
        lines = content.split('\n')
        for line in lines:
            if line.startswith('# '):
                return line[2:].strip()
            elif line.startswith('## '):
                return line[3:].strip()
        
        return "Untitled"
    
    @classmethod
    def _extract_entries_from_md(cls, content: str) -> List[Dict[str, Any]]:
        """Extract structured entries from Markdown content"""
        entries = []
        current_entry = {}
        current_section = ""
        
        lines = content.split('\n')
        
        for i, line in enumerate(lines):
            if line.startswith('# ') or line.startswith('## ') or line.startswith('### '):
                if current_entry:
                    entries.append(current_entry)
                    current_entry = {}
                
                heading_level = len(line) - len(line.lstrip('#'))
                title = line.lstrip('#').strip()
                
                if heading_level == 1:
                    current_entry["title"] = title
                    current_section = title
                elif heading_level == 2:
                    current_entry["title"] = title
                    current_entry["category"] = current_section
                elif heading_level == 3:
                    if not current_entry.get("title"):
                        current_entry["title"] = title
                    current_entry["sub_title"] = title
            
            elif line.startswith('- ') or line.startswith('* ') or line.startswith('+ '):
                bullet_content = line[2:].strip()
                if current_entry.get("content"):
                    current_entry["content"] += "\n" + bullet_content
                else:
                    current_entry["content"] = bullet_content
            
            elif line.startswith('**') and line.endswith('**'):
                bold_text = line.strip('*').strip()
                if current_entry.get("keywords"):
                    current_entry["keywords"] += ", " + bold_text
                else:
                    current_entry["keywords"] = bold_text
            
            elif line.strip() and not line.startswith(('>', '```', '|')):
                if current_entry.get("content"):
                    current_entry["content"] += "\n" + line.strip()
                else:
                    current_entry["content"] = line.strip()
        
        if current_entry:
            entries.append(current_entry)
        
        for entry in entries:
            if not entry.get("category"):
                entry["category"] = current_section if current_section else "general"
            if not entry.get("title"):
                entry["title"] = entry.get("content", "")[:50]
        
        return entries
    
    @classmethod
    def _parse_docx(cls, file_path: str) -> Dict[str, Any]:
        """Parse DOCX file"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library is not installed")
        
        try:
            doc = Document(file_path)
            title = ""
            content = ""
            entries = []
            current_entry = {}
            current_section = ""
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    if current_entry:
                        entries.append(current_entry)
                        current_entry = {}
                    continue
                
                style = para.style.name.lower()
                
                if 'heading' in style:
                    level = int(style.replace('heading', '').strip()) if 'heading' in style else 1
                    
                    if current_entry:
                        entries.append(current_entry)
                        current_entry = {}
                    
                    if level == 1:
                        title = text
                        current_section = text
                    elif level == 2:
                        current_entry["title"] = text
                        current_entry["category"] = current_section
                    elif level >= 3:
                        if not current_entry.get("title"):
                            current_entry["title"] = text
                        current_entry["sub_title"] = text
                
                else:
                    if current_entry.get("content"):
                        current_entry["content"] += "\n" + text
                    else:
                        current_entry["content"] = text
            
            if current_entry:
                entries.append(current_entry)
            
            if not title:
                title = os.path.basename(file_path).replace('.docx', '')
            
            content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            
            for entry in entries:
                if not entry.get("category"):
                    entry["category"] = current_section if current_section else "general"
                if not entry.get("title"):
                    entry["title"] = entry.get("content", "")[:50]
            
            result = {
                "title": title,
                "content": content,
                "entries": entries,
                "format": "docx",
                "file_name": os.path.basename(file_path)
            }
            
            logger.info(f"Parsed DOCX file: {title}, {len(entries)} entries")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing DOCX file: {e}")
            raise
    
    @classmethod
    def _parse_pdf(cls, file_path: str) -> Dict[str, Any]:
        """Parse PDF file"""
        if not PDF_AVAILABLE:
            raise ImportError("PyMuPDF (fitz) library is not installed")
        
        try:
            doc = fitz.open(file_path)
            content = ""
            entries = []
            current_entry = {}
            current_section = ""
            
            for page in doc:
                page_text = page.get_text()
                content += page_text + "\n"
                
                lines = page_text.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line:
                        if current_entry:
                            entries.append(current_entry)
                            current_entry = {}
                        continue
                    
                    if len(line) < 80 and line.isupper():
                        if current_entry:
                            entries.append(current_entry)
                            current_entry = {}
                        current_section = line.title()
                        current_entry["title"] = line.title()
                        current_entry["category"] = current_section
                    
                    elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '0.')):
                        if current_entry.get("content"):
                            current_entry["content"] += "\n" + line
                        else:
                            current_entry["content"] = line
                    
                    elif line.startswith(('- ', '* ', '+ ', '• ')):
                        if current_entry.get("content"):
                            current_entry["content"] += "\n" + line[2:].strip()
                        else:
                            current_entry["content"] = line[2:].strip()
                    
                    else:
                        if current_entry.get("content"):
                            current_entry["content"] += "\n" + line
                        else:
                            current_entry["content"] = line
            
            if current_entry:
                entries.append(current_entry)
            
            title = os.path.basename(file_path).replace('.pdf', '')
            
            for entry in entries:
                if not entry.get("category"):
                    entry["category"] = current_section if current_section else "general"
                if not entry.get("title"):
                    entry["title"] = entry.get("content", "")[:50]
            
            result = {
                "title": title,
                "content": content,
                "entries": entries,
                "format": "pdf",
                "file_name": os.path.basename(file_path)
            }
            
            doc.close()
            logger.info(f"Parsed PDF file: {title}, {len(entries)} entries")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing PDF file: {e}")
            raise
    
    @classmethod
    def _parse_txt(cls, file_path: str) -> Dict[str, Any]:
        """Parse TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            title = os.path.basename(file_path).replace('.txt', '')
            entries = cls._extract_entries_from_text(content)
            
            result = {
                "title": title,
                "content": content,
                "entries": entries,
                "format": "txt",
                "file_name": os.path.basename(file_path)
            }
            
            logger.info(f"Parsed TXT file: {title}, {len(entries)} entries")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing TXT file: {e}")
            raise
    
    @classmethod
    def _extract_entries_from_text(cls, content: str) -> List[Dict[str, Any]]:
        """Extract entries from plain text content"""
        entries = []
        lines = content.split('\n')
        current_entry = {}
        current_section = ""
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_entry:
                    entries.append(current_entry)
                    current_entry = {}
                continue
            
            if re.match(r'^[=]{3,}$', line):
                if current_entry:
                    entries.append(current_entry)
                    current_entry = {}
                continue
            
            if line.startswith(('## ', '### ', '# ')):
                if current_entry:
                    entries.append(current_entry)
                    current_entry = {}
                
                heading_level = len(line) - len(line.lstrip('#'))
                title = line.lstrip('#').strip()
                
                if heading_level == 1:
                    current_section = title
                    current_entry["title"] = title
                elif heading_level == 2:
                    current_entry["title"] = title
                    current_entry["category"] = current_section
            
            elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '0.')):
                if current_entry.get("content"):
                    current_entry["content"] += "\n" + line
                else:
                    current_entry["content"] = line
            
            elif line.startswith(('- ', '* ', '+ ')):
                if current_entry.get("content"):
                    current_entry["content"] += "\n" + line[2:].strip()
                else:
                    current_entry["content"] = line[2:].strip()
            
            else:
                if current_entry.get("content"):
                    current_entry["content"] += "\n" + line
                else:
                    current_entry["content"] = line
        
        if current_entry:
            entries.append(current_entry)
        
        for entry in entries:
            if not entry.get("category"):
                entry["category"] = current_section if current_section else "general"
            if not entry.get("title"):
                entry["title"] = entry.get("content", "")[:50]
        
        return entries
    
    @classmethod
    def parse_content(cls, content: str, file_name: str) -> Dict[str, Any]:
        """
        Parse content directly without file path
        
        Args:
            content: File content as string
            file_name: Original file name to determine format
            
        Returns:
            Parsed result
        """
        _, ext = os.path.splitext(file_name)
        ext = ext.lower()
        
        import tempfile
        
        with tempfile.NamedTemporaryFile(mode='w', suffix=ext, delete=False, encoding='utf-8') as f:
            f.write(content)
            temp_path = f.name
        
        try:
            return cls.parse_file(temp_path)
        finally:
            os.unlink(temp_path)